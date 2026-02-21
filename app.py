import streamlit as st
import json
import re
import os
import io
import requests
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Configuration ---
st.set_page_config(page_title="🛒 Liste de courses", page_icon="🛒", layout="wide")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(BASE_DIR, ".env"))

RECETTES_PATH = os.path.join(BASE_DIR, "recettes.json")
CATALOGUE_PATH = os.path.join(BASE_DIR, "catalogue.json")
NOTION_TOKEN = os.getenv("NOTION_TOKEN")
NOTION_PAGE_ID = os.getenv("NOTION_PAGE_ID")


# --- Chargement des données ---
def load_recettes():
    with open(RECETTES_PATH, "r", encoding="utf-8") as f:
        return json.load(f)["plats"]


def load_catalogue():
    with open(CATALOGUE_PATH, "r", encoding="utf-8") as f:
        return json.load(f)["rayons"]


def save_recettes(plats):
    """Sauvegarde la liste des plats dans recettes.json."""
    with open(RECETTES_PATH, "w", encoding="utf-8") as f:
        json.dump({"plats": plats}, f, ensure_ascii=False, indent=2)


def save_catalogue(rayons):
    """Sauvegarde les rayons dans catalogue.json."""
    with open(CATALOGUE_PATH, "w", encoding="utf-8") as f:
        json.dump({"rayons": rayons}, f, ensure_ascii=False, indent=2)


# --- Utilitaires ---
def parse_quantity(nom: str):
    """Extrait le nom de base et la quantité d'un ingrédient.
    Ex: 'Carottes (450g)' → ('Carottes', 450, 'g')
        'Tomates (3)' → ('Tomates', 3, '')
        'Crème fraîche' → ('Crème fraîche', None, '')
    """
    match = re.match(r"^(.+?)\s*\((\d+)\s*(g|kg|ml|L|cl)?\)$", nom)
    if match:
        base = match.group(1).strip()
        qty = int(match.group(2))
        unit = match.group(3) or ""
        return base, qty, unit
    return nom.strip(), None, ""


def merge_ingredients(ingredients_list):
    """Fusionne les ingrédients en dédoublonnant et cumulant les quantités.
    Retourne un dict: {rayon: [nom_affiché, ...]}
    """
    merged = {}
    for ing in ingredients_list:
        nom = ing["nom"]
        rayon = ing["rayon"]
        base, qty, unit = parse_quantity(nom)
        key = (base.lower(), rayon)

        if key in merged:
            if qty is not None and merged[key]["qty"] is not None:
                merged[key]["qty"] += qty
            elif qty is not None:
                merged[key]["qty"] = qty
                merged[key]["unit"] = unit
        else:
            merged[key] = {
                "rayon": rayon,
                "nom_base": base,
                "qty": qty,
                "unit": unit,
            }

    result = {}
    for key, data in merged.items():
        rayon = data["rayon"]
        if rayon not in result:
            result[rayon] = []
        if data["qty"] is not None:
            display = f"{data['nom_base']} ({data['qty']}{data['unit']})"
        else:
            display = data["nom_base"]
        result[rayon].append(display)

    return result


def get_recipe_ingredients(recettes, selected_names):
    """Récupère tous les ingrédients des recettes sélectionnées."""
    ingredients = []
    for recette in recettes:
        if recette["nom"] in selected_names:
            ingredients.extend(recette["ingredients"])
    return ingredients


def build_final_list(recipe_items_by_rayon, free_items_by_rayon):
    """Combine les ingrédients recettes et les articles libres, par rayon."""
    all_rayons = set(list(recipe_items_by_rayon.keys()) + list(free_items_by_rayon.keys()))

    rayon_order = [
        "BOULANGERIE",
        "LÉGUMES",
        "FRUITS",
        "AIL & FINES HERBES",
        "CHARCUTERIE",
        "TRAITEUR",
        "POISSONNERIE",
        "BOUCHERIE",
        "SURGELÉS",
        "FROMAGES",
        "YAOURTS",
        "PRODUITS LAITIERS",
        "ÉPICERIE SALÉE",
        "CUISINE DU MONDE",
        "ÉPICERIE SUCRÉE",
        "BOISSONS",
        "NOURRITURE BÉBÉ",
        "HYGIÈNE & DIVERS",
    ]

    final = {}
    for rayon in rayon_order:
        if rayon in all_rayons:
            items = set()
            items.update(recipe_items_by_rayon.get(rayon, []))
            items.update(free_items_by_rayon.get(rayon, []))
            if items:
                final[rayon] = sorted(items)

    for rayon in sorted(all_rayons - set(rayon_order)):
        items = set()
        items.update(recipe_items_by_rayon.get(rayon, []))
        items.update(free_items_by_rayon.get(rayon, []))
        if items:
            final[rayon] = sorted(items)

    return final


def export_to_notion(final_list, selected_recipes):
    """Crée une page Notion avec des cases à cocher via l'API."""
    if not NOTION_TOKEN or not NOTION_PAGE_ID:
        return False, "Configuration Notion manquante. Vérifiez le fichier .env.", None

    headers = {
        "Authorization": f"Bearer {NOTION_TOKEN}",
        "Content-Type": "application/json",
        "Notion-Version": "2022-06-28",
    }

    date_str = datetime.now().strftime("%d/%m/%Y")
    title = f"🛒 Liste de courses — {date_str}"

    children = []

    if selected_recipes:
        children.append({
            "object": "block",
            "type": "paragraph",
            "paragraph": {
                "rich_text": [{
                    "type": "text",
                    "text": {"content": f"🍽️ {' • '.join(selected_recipes)}"},
                    "annotations": {"italic": True, "color": "gray"},
                }]
            }
        })
        children.append({"object": "block", "type": "divider", "divider": {}})

    for rayon, items in final_list.items():
        children.append({
            "object": "block",
            "type": "heading_2",
            "heading_2": {
                "rich_text": [{"type": "text", "text": {"content": rayon}}]
            }
        })
        for item in items:
            children.append({
                "object": "block",
                "type": "to_do",
                "to_do": {
                    "rich_text": [{"type": "text", "text": {"content": item}}],
                    "checked": False,
                }
            })

    payload = {
        "parent": {"page_id": NOTION_PAGE_ID},
        "properties": {
            "title": [{"text": {"content": title}}]
        },
        "children": children[:100],
    }

    try:
        resp = requests.post(
            "https://api.notion.com/v1/pages",
            headers=headers,
            json=payload,
            timeout=15,
        )

        if resp.status_code == 200:
            page_url = resp.json().get("url", "")

            if len(children) > 100:
                page_id = resp.json()["id"]
                for i in range(100, len(children), 100):
                    batch = children[i:i+100]
                    requests.patch(
                        f"https://api.notion.com/v1/blocks/{page_id}/children",
                        headers=headers,
                        json={"children": batch},
                        timeout=15,
                    )

            return True, "Page créée dans Notion !", page_url
        else:
            error = resp.json().get("message", resp.text)
            return False, f"Erreur Notion : {error}", None

    except requests.exceptions.Timeout:
        return False, "Timeout : Notion n'a pas répondu.", None
    except Exception as e:
        return False, f"Erreur : {str(e)}", None


def export_to_docx(final_list, selected_recipes):
    """Génère un fichier Word de la liste de courses."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("Liste de courses", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Semaine du {datetime.now().strftime('%d/%m/%Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    if selected_recipes:
        doc.add_paragraph()
        plats_para = doc.add_paragraph()
        plats_run = plats_para.add_run("Plats : ")
        plats_run.bold = True
        plats_run.font.size = Pt(10)
        plats_text = plats_para.add_run(" • ".join(selected_recipes))
        plats_text.font.size = Pt(10)
        plats_text.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    for rayon, items in final_list.items():
        heading = doc.add_heading(rayon, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(46, 117, 182)
            run.font.size = Pt(13)

        for item in items:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(item)
            run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def add_ingredient_to_catalogue(catalogue, nom_ingredient, rayon_nom):
    """Ajoute un ingrédient au catalogue s'il n'y est pas déjà.
    Retourne True si le catalogue a été modifié."""
    for rayon in catalogue:
        if rayon["nom"] == rayon_nom:
            # Vérifier si l'article existe déjà (insensible à la casse)
            existing_lower = [a.lower() for a in rayon["articles"]]
            if nom_ingredient.lower() not in existing_lower:
                rayon["articles"].append(nom_ingredient)
                rayon["articles"].sort(key=str.lower)
                return True
            return False
    # Rayon inexistant → le créer
    catalogue.append({"nom": rayon_nom, "articles": [nom_ingredient]})
    return True


# --- Chargement ---
recettes = load_recettes()
catalogue = load_catalogue()

# --- Session state ---
if "checked_items" not in st.session_state:
    st.session_state.checked_items = set()
if "new_recipe_ingredients" not in st.session_state:
    st.session_state.new_recipe_ingredients = []


# --- Interface ---
st.title("🛒 Liste de courses")

tab_recettes, tab_produits, tab_liste, tab_nouvelle = st.tabs(
    ["🍽️ Recettes", "🏪 Produits", "📋 Ma liste", "➕ Nouvelle recette"]
)

# =====================
# ONGLET 1 : RECETTES
# =====================
with tab_recettes:
    st.header("Sélectionnez vos plats de la semaine")

    cols = st.columns(2)
    for i, recette in enumerate(recettes):
        with cols[i % 2]:
            ingredients_str = ", ".join(ing["nom"] for ing in recette["ingredients"])
            st.checkbox(
                recette["nom"],
                key=f"recette_{i}",
                help=ingredients_str,
            )

    _selected = [r["nom"] for i, r in enumerate(recettes) if st.session_state.get(f"recette_{i}", False)]
    if _selected:
        st.divider()
        st.subheader("Ingrédients sélectionnés")
        _ingredients = get_recipe_ingredients(recettes, _selected)
        _by_rayon = merge_ingredients(_ingredients)
        for rayon, items in sorted(_by_rayon.items()):
            st.markdown(f"**{rayon}**")
            for item in items:
                st.markdown(f"- {item}")

# =====================
# ONGLET 2 : PRODUITS
# =====================
with tab_produits:
    st.header("Ajoutez des articles par rayon")

    for rayon in catalogue:
        with st.expander(f"🏷️ {rayon['nom']} ({len(rayon['articles'])} articles)"):
            for j, article in enumerate(rayon["articles"]):
                st.checkbox(
                    article,
                    key=f"cat_{rayon['nom']}_{j}",
                )

# ==============================
# ONGLET 4 : NOUVELLE RECETTE
# ==============================
with tab_nouvelle:
    st.header("Ajouter une nouvelle recette")

    # Nom de la recette
    new_recipe_name = st.text_input(
        "Nom de la recette",
        placeholder="Ex : Gratin dauphinois",
        key="new_recipe_name",
    )

    st.divider()
    st.subheader("Ingrédients")

    # Liste des rayons existants
    rayon_names = [r["nom"] for r in catalogue]

    # Formulaire d'ajout d'ingrédient
    with st.form("add_ingredient_form", clear_on_submit=True):
        col_ing, col_rayon = st.columns([2, 1])
        with col_ing:
            ing_name = st.text_input(
                "Nom de l'ingrédient",
                placeholder="Ex : Pommes de terre (500g)",
            )
        with col_rayon:
            ing_rayon = st.selectbox(
                "Rayon",
                options=rayon_names,
                index=0,
            )
        submitted = st.form_submit_button("➕ Ajouter l'ingrédient")

        if submitted and ing_name.strip():
            st.session_state.new_recipe_ingredients.append({
                "nom": ing_name.strip(),
                "rayon": ing_rayon,
            })

    # Afficher les ingrédients ajoutés
    if st.session_state.new_recipe_ingredients:
        st.markdown("---")
        st.markdown("**Ingrédients ajoutés :**")
        to_remove = None
        for idx, ing in enumerate(st.session_state.new_recipe_ingredients):
            col_display, col_del = st.columns([4, 1])
            with col_display:
                st.markdown(f"- **{ing['nom']}** — _{ing['rayon']}_")
            with col_del:
                if st.button("🗑️", key=f"del_ing_{idx}"):
                    to_remove = idx

        if to_remove is not None:
            st.session_state.new_recipe_ingredients.pop(to_remove)
            st.rerun()

        st.markdown("---")

        # Bouton de sauvegarde
        if st.button("💾 Enregistrer la recette", type="primary"):
            if not new_recipe_name.strip():
                st.error("Donnez un nom à la recette.")
            else:
                # Vérifier que le nom n'existe pas déjà
                existing_names = [r["nom"].lower() for r in recettes]
                if new_recipe_name.strip().lower() in existing_names:
                    st.error(f"La recette « {new_recipe_name.strip()} » existe déjà.")
                else:
                    # Ajouter la recette
                    new_recipe = {
                        "nom": new_recipe_name.strip(),
                        "ingredients": list(st.session_state.new_recipe_ingredients),
                    }
                    recettes.append(new_recipe)
                    save_recettes(recettes)

                    # Mettre à jour le catalogue avec les nouveaux ingrédients
                    catalogue_modified = False
                    for ing in st.session_state.new_recipe_ingredients:
                        if add_ingredient_to_catalogue(catalogue, ing["nom"], ing["rayon"]):
                            catalogue_modified = True
                    if catalogue_modified:
                        save_catalogue(catalogue)

                    # Réinitialiser le formulaire
                    st.session_state.new_recipe_ingredients = []
                    st.success(f"✅ Recette « {new_recipe_name.strip()} » enregistrée avec {len(new_recipe['ingredients'])} ingrédient(s) !")
                    st.balloons()
                    st.rerun()
    else:
        st.caption("Ajoutez des ingrédients via le formulaire ci-dessus.")

# ============================================
# CALCUL DE LA LISTE FINALE (hors des tabs)
# ============================================
selected_recipes_final = []
for i, recette in enumerate(recettes):
    if st.session_state.get(f"recette_{i}", False):
        selected_recipes_final.append(recette["nom"])

recipe_ingredients_final = get_recipe_ingredients(recettes, selected_recipes_final)
recipe_by_rayon_final = merge_ingredients(recipe_ingredients_final)

free_items_final = {}
for rayon in catalogue:
    items = []
    for j, article in enumerate(rayon["articles"]):
        if st.session_state.get(f"cat_{rayon['nom']}_{j}", False):
            items.append(article)
    if items:
        free_items_final[rayon["nom"]] = items

final_list = build_final_list(recipe_by_rayon_final, free_items_final)

# =====================
# ONGLET 3 : MA LISTE
# =====================
with tab_liste:

    if final_list:
        st.header("📋 Ma liste de courses")

        if selected_recipes_final:
            st.caption(f"🍽️ Plats : {' • '.join(selected_recipes_final)}")

        st.divider()

        # Compteur
        total = sum(len(items) for items in final_list.values())
        checked_count = len(
            [
                item
                for rayon, items in final_list.items()
                for item in items
                if f"check_{rayon}_{item}" in st.session_state.checked_items
            ]
        )
        st.progress(
            checked_count / total if total > 0 else 0,
            text=f"✅ {checked_count}/{total} articles",
        )

        # Liste avec cases à cocher
        for rayon, items in final_list.items():
            st.subheader(rayon)
            for item in items:
                check_key = f"check_{rayon}_{item}"
                checked = st.checkbox(
                    item,
                    key=check_key,
                    value=check_key in st.session_state.checked_items,
                )
                if checked:
                    st.session_state.checked_items.add(check_key)
                elif check_key in st.session_state.checked_items:
                    st.session_state.checked_items.discard(check_key)

        st.divider()

        # Boutons d'action
        col1, col2, col3 = st.columns(3)
        with col1:
            docx_buffer = export_to_docx(final_list, selected_recipes_final)
            st.download_button(
                label="📥 Exporter en Word",
                data=docx_buffer,
                file_name=f"Liste_courses_{datetime.now().strftime('%Y-%m-%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col2:
            if st.button("📝 Envoyer vers Notion"):
                with st.spinner("Création de la page Notion..."):
                    success, message, url = export_to_notion(final_list, selected_recipes_final)
                if success:
                    st.success(message)
                    if url:
                        st.markdown(f"[🔗 Ouvrir dans Notion]({url})")
                else:
                    st.error(message)
        with col3:
            if st.button("🗑️ Réinitialiser les coches"):
                st.session_state.checked_items = set()
                st.rerun()
    else:
        st.info(
            "👈 Sélectionnez des recettes dans l'onglet **Recettes** "
            "ou ajoutez des articles depuis les **Produits** pour constituer votre liste."
        )
